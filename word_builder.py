from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_word_report(
    session: Dict[str, Any],
    data_root: Path
) -> None:

    inspection_id = session["inspection_id"]
    template_path = data_root / "templates" / "report_template.docx"
    output_dir = data_root / "outputs" / inspection_id
    output_path = output_dir / "report.docx"

    # Load template
    document = Document(template_path)

    document.add_paragraph("Construction Observations")

    # ===== Observations =====
    observations = session.get("observations", [])

    report_obs = [
        obs for obs in observations
        if obs.get("include_in_report", True)
    ]

    report_obs.sort(key=lambda x: x.get("number", 0))

    for obs in report_obs:
        number = obs.get("number")
        raw_text = (obs.get("raw_text") or "").strip()

        p = document.add_paragraph(raw_text, style="List Paragraph")
        if raw_text:
            p.add_run(raw_text)
        else:
            p.add_run("")

        photos = obs.get("photos", [])

        if photos:
            photo_line = document.add_paragraph()
            photo_line.add_run(
                f"Photo(s): {', '.join(photos)}"
            )

    # ===== Photo Appendix =====
    document.add_paragraph("Photo Appendix")

    photo_ids = []

    for obs in report_obs:
        for ph in obs.get("photos", []):
            if ph not in photo_ids:
                photo_ids.append(ph)

    for ph in photo_ids:
        document.add_paragraph(ph)

        image_path = (
            data_root
            / "tmp_photos"
            / inspection_id
            / f"{ph}.jpg"
        )

        if image_path.exists():
            pic = document.add_picture(
                str(image_path),
                width=Cm(10)
            )
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(output_path)