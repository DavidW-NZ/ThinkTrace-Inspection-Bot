"""template_word_builder.py

Template-driven Word report builder (Strict Mode) – v1.3 (patched)

Locked behaviours:
- Users control layout by placing placeholders in their own .docx template.
- Mandatory blocks (must exist exactly once):
    {OBSERVATIONS_BLOCK}
    {PHOTO_APPENDIX_BLOCK}
- Optional block (0 or 1):
    {HIDDEN_ITEMS_BLOCK}
- Review items are included in the SAME block as observations ({OBSERVATIONS_BLOCK}).
- Main report numbering is continuous and excludes hidden items:
    - Observations shown = include_in_report == True
    - Review items numbering continues after shown observations
- Hidden block (if present) outputs ALL hidden items and shows their photos inline under each item.
- Photo appendix includes ONLY non-hidden photos.
- Images inserted are 10cm wide, centred, inline.

Patch goals:
- Fix missing helper functions (_paragraph_contains_placeholder, _unique_paragraph_key, etc.).
- Handle Word splitting placeholders across multiple runs / XML text nodes.
  (We do paragraph-local merge ONLY when a placeholder is detected.)
- Support additional single-value placeholders:
    {CREATED_DATE}, {CREATED_TIME}, {CONFIRMED_DATE}, {CONFIRMED_TIME}
"""

from __future__ import annotations

import re
from datetime import datetime, timezone
from zoneinfo import ZoneInfo
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.text.paragraph import Paragraph


# ---------------------------
# Placeholders
# ---------------------------

_PH_OBS_BLOCK = "{OBSERVATIONS_BLOCK}"
_PH_PHOTO_APPENDIX_BLOCK = "{PHOTO_APPENDIX_BLOCK}"
_PH_HIDDEN_ITEMS_BLOCK = "{HIDDEN_ITEMS_BLOCK}"

_PH_ACTION_REQUIRED_BLOCK = "{ACTION_REQUIRED_BLOCK}"
_PH_ACTION_COMPLETED_BLOCK = "{ACTION_COMPLETED_BLOCK}"
_PH_ACTION_REQUIRED_PHOTOS_BLOCK = "{ACTION_REQUIRED_PHOTOS_BLOCK}"
_PH_ACTION_COMPLETED_PHOTOS_BLOCK = "{ACTION_COMPLETED_PHOTOS_BLOCK}"

_HIDDEN_NOTE = "[HIDDEN – excluded from report]"

_PH_RE = re.compile(r"^PH-\d{3}$")

LOCAL_TZ = ZoneInfo("Pacific/Auckland")

# ---------------------------
# Public API
# ---------------------------

def build_word_report_from_template(session: Dict[str, Any], data_root: Path) -> None:
    """Build outputs/<inspection_id>/report.docx by filling templates/report_template.docx."""

    inspection_id = _require_str(session, "inspection_id")
    output_dir = data_root / "outputs" / inspection_id
    output_dir.mkdir(parents=True, exist_ok=True)

    template_path = data_root / "templates" / "report_template.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template_path))

    # 1) Single-value replacements (strict run replace + paragraph-local fallback)
    single_values = _build_single_values(session)
    _replace_single_value_placeholders(doc, single_values)

    # 2) Block replacements
    observations = session.get("observations", []) or []
    review_items = session.get("review_items", []) or []

    _replace_block_observations(doc, observations, review_items)
    _replace_block_photo_appendix(doc, session, observations, data_root)

    # Action blocks (optional)
    actions_required = session.get("actions_required", []) or []
    actions_completed = session.get("actions_completed", []) or []

    _replace_block_action_required_if_present(doc, actions_required)
    _replace_block_action_completed_if_present(doc, actions_completed)
    _replace_block_action_required_photos_if_present(doc, session, actions_required, data_root)
    _replace_block_action_completed_photos_if_present(doc, session, actions_completed, data_root)

    # Optional hidden block
    _replace_block_hidden_items_if_present(doc, session, observations, data_root)

    out_path = output_dir / "report.docx"
    doc.save(str(out_path))


# ---------------------------
# Single-value placeholders
# ---------------------------

def _build_single_values(session: Dict[str, Any]) -> Dict[str, str]:
    header = session.get("header", {}) or {}

    project_id = str(session.get("project_id", "") or "")
    inspection_id = str(session.get("inspection_id", "") or "")

    created_iso = str(session.get("created_at", "") or "")
    confirmed_iso = str(session.get("confirmed_at", "") or "")

    created_dt = _parse_dt(created_iso)
    confirmed_dt = _parse_dt(confirmed_iso)

    if created_dt:
        created_dt = _to_nz(created_dt)
    if confirmed_dt:
        confirmed_dt = _to_nz(confirmed_dt)

    created_date = created_dt.strftime("%d-%m-%Y") if created_dt else ""
    created_time = created_dt.strftime("%H:%M") if created_dt else ""
    confirmed_date = confirmed_dt.strftime("%d-%m-%Y") if confirmed_dt else ""
    confirmed_time = confirmed_dt.strftime("%H:%M") if confirmed_dt else ""

    created_at = (
        f"{created_date} {created_time}".strip()
        if (created_date or created_time)
        else _format_dt_local(created_iso)
    )
    confirmed_at = (
        f"{confirmed_date} {confirmed_time}".strip()
        if (confirmed_date or confirmed_time)
        else _format_dt_local(confirmed_iso)
    )

    location = str(header.get("location_text", "") or "")
    weather = "" if header.get("weather", None) is None else str(header.get("weather", "") or "")
    cor_no = str(session.get("cor_no", "") or "")
    inspection_title = str(header.get("title", "") or "")

    return {
        "{PROJECT_ID}": project_id,
        "{INSPECTION_ID}": inspection_id,

        "{CREATED_AT}": created_at,
        "{CREATED_DATE}": created_date,
        "{CREATED_TIME}": created_time,

        "{CONFIRMED_AT}": confirmed_at,
        "{CONFIRMED_DATE}": confirmed_date,
        "{CONFIRMED_TIME}": confirmed_time,

        "{LOCATION}": location,
        "{WEATHER}": weather,
        "{COR_NO}": cor_no,
        "{INSPECTION_TITLE}": inspection_title,
    }


def _replace_single_value_placeholders(doc: Document, mapping: Dict[str, str]) -> None:
    """Replace single-value placeholders.

    Pass 1: run-level replacement (fast, preserves formatting).
    Pass 2 (fallback): if a placeholder is split across runs/text nodes, do a paragraph-local merge,
    replace, then keep formatting of the first run.
    """

    def _run_level(paragraph: Any) -> None:
        if not getattr(paragraph, "runs", None):
            return
        for run in paragraph.runs:
            if not run.text:
                continue
            for ph, value in mapping.items():
                if ph in run.text:
                    run.text = run.text.replace(ph, value)

    def _paragraph_fallback(paragraph: Any) -> None:
        if not getattr(paragraph, "runs", None):
            return
        full_text = "".join([(r.text or "") for r in paragraph.runs])
        if not full_text:
            return

        replaced = full_text
        changed = False

        # direct
        for ph, value in mapping.items():
            if ph in replaced:
                replaced = replaced.replace(ph, value)
                changed = True

        # tolerant (handles split runs, zero-width, and incidental spaces)
        if not changed:
            for ph, value in mapping.items():
                rgx = _placeholder_tolerant_regex(ph)
                if rgx.search(replaced):
                    replaced = rgx.sub(value, replaced)
                    changed = True

        if changed:
            paragraph.runs[0].text = replaced
            for r in paragraph.runs[1:]:
                r.text = ""

    # Body + tables
    for paragraph in _iter_all_paragraphs(doc):
        _run_level(paragraph)
        _paragraph_fallback(paragraph)

    # Headers/footers (common for IDs/dates)
    for section in doc.sections:
        for paragraph in _iter_header_footer_paragraphs(section):
            _run_level(paragraph)
            _paragraph_fallback(paragraph)


# ---------------------------
# Block placeholders
# ---------------------------

def _replace_block_observations(
    doc: Document,
    observations: Sequence[Dict[str, Any]],
    review_items: Sequence[Dict[str, Any]],
) -> None:
    """Replace {OBSERVATIONS_BLOCK} with:
    - Non-hidden observations (include_in_report=True), numbered 1..N
    - Then review items, numbered N+1..
    """

    host_para = _find_single_block_paragraph(doc, _PH_OBS_BLOCK, required=True)
    insert_after = host_para

    report_obs = [o for o in observations if bool(o.get("include_in_report", True))]
    report_obs.sort(key=lambda x: int(x.get("number", 0) or 0))

    display_no = 0
    for obs in report_obs:
        display_no += 1

        # Prefer AI rewrite if present, else fallback raw_text
        text = _normalize_text(str(obs.get("rewritten_text") or "") or str(obs.get("raw_text", "") or ""))

        insert_after = _add_paragraph_after(insert_after, f"{display_no}. {text}".rstrip())

        ph_list = _validate_ph_list(obs.get("photos", []) or [])
        if ph_list:
            insert_after = _add_paragraph_after(insert_after, f"Photo(s): {', '.join(ph_list)}")

        insert_after = _add_paragraph_after(insert_after, "")

    sorted_review = list(review_items or [])
    sorted_review.sort(key=lambda x: int(x.get("number", 0) or 0))

    for item in sorted_review:
        display_no += 1
        # Prefer AI rewrite if present on review item
        text = _normalize_text(str(item.get("rewritten_text") or "") or str(item.get("text", "") or ""))
        insert_after = _add_paragraph_after(insert_after, f"{display_no}. {text}".rstrip())
        insert_after = _add_paragraph_after(insert_after, "")

    _delete_paragraph(host_para)


def _replace_block_photo_appendix(
    doc: Document,
    session: Dict[str, Any],
    observations: Sequence[Dict[str, Any]],
    data_root: Path,
) -> None:
    """Replace {PHOTO_APPENDIX_BLOCK} with ONLY non-hidden photos."""

    host_para = _find_single_block_paragraph(doc, _PH_PHOTO_APPENDIX_BLOCK, required=True)
    inspection_id = _require_str(session, "inspection_id")

    report_obs = [o for o in observations if bool(o.get("include_in_report", True))]

    photo_ids: List[str] = []
    seen = set()
    for obs in report_obs:
        for ph in _validate_ph_list(obs.get("photos", []) or []):
            if ph not in seen:
                seen.add(ph)
                photo_ids.append(ph)

    photo_ids.sort(key=_ph_sort_key)

    insert_after = host_para
    photo_dir = data_root / "tmp_photos" / inspection_id

    for ph in photo_ids:
        insert_after = _add_paragraph_after(insert_after, ph)
        img_path = _find_photo_file(photo_dir, ph)
        if img_path is not None:
            insert_after = _add_picture_after(insert_after, img_path)
        insert_after = _add_paragraph_after(insert_after, "")

    _delete_paragraph(host_para)



def _replace_block_action_required_if_present(
    doc: Document,
    actions_required: Sequence[Dict[str, Any]],
) -> None:
    host_para = _find_single_block_paragraph(doc, _PH_ACTION_REQUIRED_BLOCK, required=False)
    if host_para is None:
        return

    insert_after = host_para
    items = list(actions_required or [])
    items.sort(key=lambda x: int(x.get("number", 0) or 0))

    for it in items:
        n = int(it.get("number", 0) or 0)
        text = _normalize_text(str((it.get("rewritten_text") or it.get("raw_text") or "")))
        insert_after = _add_paragraph_after(insert_after, f"{n}. {text}".rstrip())

        ph_list = _validate_ph_list(it.get("photos", []) or [])
        if ph_list:
            insert_after = _add_paragraph_after(insert_after, f"Photo(s): {', '.join(ph_list)}")

        insert_after = _add_paragraph_after(insert_after, "")

    _delete_paragraph(host_para)


def _replace_block_action_completed_if_present(
    doc: Document,
    actions_completed: Sequence[Dict[str, Any]],
) -> None:
    host_para = _find_single_block_paragraph(doc, _PH_ACTION_COMPLETED_BLOCK, required=False)
    if host_para is None:
        return

    insert_after = host_para
    items = list(actions_completed or [])
    items.sort(key=lambda x: int(x.get("number", 0) or 0))

    for it in items:
        n = int(it.get("number", 0) or 0)
        text = _normalize_text(str((it.get("rewritten_text") or it.get("raw_text") or "")))
        insert_after = _add_paragraph_after(insert_after, f"{n}. {text}".rstrip())

        ph_list = _validate_ph_list(it.get("photos", []) or [])
        if ph_list:
            insert_after = _add_paragraph_after(insert_after, f"Photo(s): {', '.join(ph_list)}")

        insert_after = _add_paragraph_after(insert_after, "")

    _delete_paragraph(host_para)


def _replace_block_action_required_photos_if_present(
    doc: Document,
    session: Dict[str, Any],
    actions_required: Sequence[Dict[str, Any]],
    data_root: Path,
) -> None:
    host_para = _find_single_block_paragraph(doc, _PH_ACTION_REQUIRED_PHOTOS_BLOCK, required=False)
    if host_para is None:
        return

    inspection_id = _require_str(session, "inspection_id")
    photo_dir = data_root / "tmp_photos" / inspection_id

    photo_ids: List[str] = []
    seen = set()
    for it in actions_required or []:
        for ph in _validate_ph_list(it.get("photos", []) or []):
            if ph not in seen:
                seen.add(ph)
                photo_ids.append(ph)

    photo_ids.sort(key=_ph_sort_key)

    insert_after = host_para
    for ph in photo_ids:
        insert_after = _add_paragraph_after(insert_after, ph)
        img_path = _find_photo_file(photo_dir, ph)
        if img_path is not None:
            insert_after = _add_picture_after(insert_after, img_path)
        insert_after = _add_paragraph_after(insert_after, "")

    _delete_paragraph(host_para)


def _replace_block_action_completed_photos_if_present(
    doc: Document,
    session: Dict[str, Any],
    actions_completed: Sequence[Dict[str, Any]],
    data_root: Path,
) -> None:
    host_para = _find_single_block_paragraph(doc, _PH_ACTION_COMPLETED_PHOTOS_BLOCK, required=False)
    if host_para is None:
        return

    inspection_id = _require_str(session, "inspection_id")
    photo_dir = data_root / "tmp_photos" / inspection_id

    photo_ids: List[str] = []
    seen = set()
    for it in actions_completed or []:
        for ph in _validate_ph_list(it.get("photos", []) or []):
            if ph not in seen:
                seen.add(ph)
                photo_ids.append(ph)

    photo_ids.sort(key=_ph_sort_key)

    insert_after = host_para
    for ph in photo_ids:
        insert_after = _add_paragraph_after(insert_after, ph)
        img_path = _find_photo_file(photo_dir, ph)
        if img_path is not None:
            insert_after = _add_picture_after(insert_after, img_path)
        insert_after = _add_paragraph_after(insert_after, "")

    _delete_paragraph(host_para)


def _replace_block_hidden_items_if_present(
    doc: Document,
    session: Dict[str, Any],
    observations: Sequence[Dict[str, Any]],
    data_root: Path,
) -> None:
    """Replace {HIDDEN_ITEMS_BLOCK} with ALL hidden observations (include_in_report=False)."""

    host_para = _find_single_block_paragraph(doc, _PH_HIDDEN_ITEMS_BLOCK, required=False)
    if host_para is None:
        return

    inspection_id = _require_str(session, "inspection_id")
    photo_dir = data_root / "tmp_photos" / inspection_id

    hidden_obs = [o for o in observations if not bool(o.get("include_in_report", True))]
    hidden_obs.sort(key=lambda x: int(x.get("number", 0) or 0))

    insert_after = host_para

    for obs in hidden_obs:
        n = int(obs.get("number", 0) or 0)  # traceability
        text = _normalize_text(str(obs.get("raw_text", "") or ""))
        insert_after = _add_paragraph_after(insert_after, f"{n}. {text}".rstrip())

        ph_list = _validate_ph_list(obs.get("photos", []) or [])
        if ph_list:
            insert_after = _add_paragraph_after(insert_after, f"Photo(s): {', '.join(ph_list)}")

        for ph in ph_list:
            img_path = _find_photo_file(photo_dir, ph)
            if img_path is not None:
                insert_after = _add_picture_after(insert_after, img_path)

        insert_after = _add_paragraph_after(insert_after, _HIDDEN_NOTE)
        insert_after = _add_paragraph_after(insert_after, "")

    _delete_paragraph(host_para)


# ---------------------------
# Helpers: paragraph iteration
# ---------------------------

def _iter_all_paragraphs(doc: Document) -> Iterable[Any]:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Any) -> Iterable[Any]:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for t in cell.tables:
                yield from _iter_table_paragraphs(t)


def _iter_header_footer_paragraphs(section: Any) -> Iterable[Any]:
    for p in section.header.paragraphs:
        yield p
    for t in section.header.tables:
        yield from _iter_table_paragraphs(t)

    for p in section.footer.paragraphs:
        yield p
    for t in section.footer.tables:
        yield from _iter_table_paragraphs(t)


# ---------------------------
# Helpers: block locate + insert + delete
# ---------------------------

def _placeholder_tolerant_regex(placeholder: str) -> re.Pattern:
    # Allow whitespace and common zero-width chars between each character.
    gap = r"(?:\s|\u200b|\ufeff|\u200c|\u200d)*"
    escaped = [re.escape(ch) for ch in placeholder]
    expr = gap.join(escaped)
    return re.compile(expr)


def _paragraph_full_text(paragraph: Any) -> str:
    if not getattr(paragraph, "runs", None):
        return ""
    return "".join([(r.text or "") for r in paragraph.runs])


def _paragraph_contains_placeholder(paragraph: Any, placeholder: str) -> bool:
    text = _paragraph_full_text(paragraph)
    if not text:
        return False
    if placeholder in text:
        return True
    return bool(_placeholder_tolerant_regex(placeholder).search(text))


def _paragraph_is_pure_placeholder(paragraph: Any, placeholder: str) -> bool:
    text = _paragraph_full_text(paragraph)
    if not text:
        return False

    # Remove common zero-width chars then strip.
    cleaned = (
        text.replace("\u200b", "")
        .replace("\ufeff", "")
        .replace("\u200c", "")
        .replace("\u200d", "")
        .strip()
    )

    if cleaned == placeholder:
        return True

    # tolerant equality: entire cleaned string matches tolerant placeholder regex.
    rgx = _placeholder_tolerant_regex(placeholder)
    return bool(rgx.fullmatch(cleaned))


def _unique_paragraph_key(paragraph: Any) -> str:
    # Use XML element id for stable de-dup across multiple iterators.
    return str(id(paragraph._p))


def _find_single_block_paragraph(doc: Document, placeholder: str, required: bool) -> Optional[Any]:
    matches_map: Dict[str, Any] = {}

    for p in _iter_all_paragraphs(doc):
        if _paragraph_contains_placeholder(p, placeholder):
            matches_map[_unique_paragraph_key(p)] = p

    for section in doc.sections:
        for p in _iter_header_footer_paragraphs(section):
            if _paragraph_contains_placeholder(p, placeholder):
                matches_map[_unique_paragraph_key(p)] = p

    matches: List[Any] = list(matches_map.values())

    if not matches:
        if required:
            raise ValueError(f"Missing mandatory placeholder: {placeholder}")
        return None

    if len(matches) > 1:
        raise ValueError(f"Placeholder appears multiple times (must be unique): {placeholder}")

    host = matches[0]
    if not _paragraph_is_pure_placeholder(host, placeholder):
        # Avoid deleting other user text in the same paragraph.
        raise ValueError(
            f"Placeholder must be the only content in its paragraph: {placeholder}"
        )

    return host


def _add_paragraph_after(paragraph: Any, text: str) -> Any:
    """Insert a new paragraph after `paragraph` with `text`."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def _add_picture_after(paragraph: Any, image_path: Path) -> Any:
    new_para = _add_paragraph_after(paragraph, "")
    run = new_para.add_run()
    run.add_picture(str(image_path), width=Cm(10))
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return new_para


def _delete_paragraph(paragraph: Any) -> None:
    p = paragraph._p
    p.getparent().remove(p)


# ---------------------------
# Helpers: validation + formatting
# ---------------------------

def _validate_ph_list(value: Any) -> List[str]:
    if value is None:
        return []
    if not isinstance(value, list):
        raise TypeError("photos must be a list of strings")

    out: List[str] = []
    for item in value:
        if not isinstance(item, str):
            raise TypeError("Photo ID must be string")
        if not _PH_RE.match(item):
            raise ValueError(f"Invalid photo ID: {item}")
        out.append(item)
    return out


def _ph_sort_key(ph: str) -> int:
    try:
        return int(ph.split("-")[1])
    except Exception:
        return 10**9


def _find_photo_file(photo_dir: Path, ph: str) -> Optional[Path]:
    candidates = [
        photo_dir / f"{ph}.jpg",
        photo_dir / f"{ph}.jpeg",
        photo_dir / f"{ph}.png",
        photo_dir / f"{ph}.JPG",
        photo_dir / f"{ph}.JPEG",
        photo_dir / f"{ph}.PNG",
    ]
    for c in candidates:
        if c.exists():
            return c
    return None


def _normalize_text(s: str) -> str:
    lines = (s or "").strip().splitlines()
    return "\n".join(line.rstrip() for line in lines)


def _to_nz(dt: datetime) -> datetime:
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(LOCAL_TZ)


def _parse_dt(dt_str: str) -> Optional[datetime]:
    if not dt_str:
        return None
    s = str(dt_str).strip()
    if not s:
        return None
    try:
        if s.endswith("Z"):
            s = s[:-1] + "+00:00"
        return datetime.fromisoformat(s)
    except Exception:
        return None


def _format_dt_local(dt_str: str) -> str:
    if not dt_str:
        return ""
    try:
        s = str(dt_str).strip()
        if s.endswith("Z"):
            s = s[:-1] + "+00:00"
        dt = datetime.fromisoformat(s)
        return _to_nz(dt).strftime("%d-%m-%Y %H:%M")
    except Exception:
        return ""


def _require_str(d: Dict[str, Any], key: str) -> str:
    v = str(d.get(key, "") or "")
    if not v:
        raise KeyError(f"Missing required field: {key}")
    return v
